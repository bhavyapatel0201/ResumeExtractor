#!/usr/bin/env python3
"""Inspect the SAMPLE final.docx template structure."""

try:
    from docx import Document
    
    template_path = r'SAMPLE FINAL\SAMPLE final.docx'
    doc = Document(template_path)
    
    print("=" * 80)
    print("TEMPLATE STRUCTURE ANALYSIS")
    print("=" * 80)
    
    print(f"\nTotal Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Tables: {len(doc.tables)}")
    print(f"Total Sections: {len(doc.sections)}")
    
    print("\n" + "=" * 80)
    print("PARAGRAPHS CONTENT:")
    print("=" * 80)
    for i, para in enumerate(doc.paragraphs[:30]):
        style = para.style.name if para.style else "None"
        text = para.text[:100] if para.text else "(empty)"
        print(f"{i:2d} | Style: {style:20s} | Text: {text}")
    
    print("\n" + "=" * 80)
    print("TABLES INFO:")
    print("=" * 80)
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i}:")
        print(f"  Rows: {len(table.rows)}, Columns: {len(table.columns)}")
        if table.rows:
            for j, row in enumerate(table.rows[:3]):
                row_text = " | ".join([cell.text[:30] for cell in row.cells])
                print(f"  Row {j}: {row_text}")
    
    print("\n" + "=" * 80)
    print("DOCUMENT PROPERTIES:")
    print("=" * 80)
    print(f"Core Properties: {doc.core_properties}")
    
except ImportError:
    print("ERROR: python-docx is not installed")
    print("Install with: pip install python-docx")
except FileNotFoundError:
    print(f"ERROR: Template file not found: {template_path}")
except Exception as e:
    print(f"ERROR: {e}")
