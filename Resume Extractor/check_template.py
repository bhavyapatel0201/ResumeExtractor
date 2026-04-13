#!/usr/bin/env python3
from docx import Document

doc = Document(r'SAMPLE FINAL\SAMPLE final.docx')
print('=== TEMPLATE STRUCTURE ===')
print(f'Total Paragraphs: {len(doc.paragraphs)}')
print(f'Total Tables: {len(doc.tables)}')
print(f'Total Sections: {len(doc.sections)}')
print()
print('=== SECTION DETAILS ===')
for i, section in enumerate(doc.sections):
    print(f'Section {i}:')
    print(f'  - Has Header: {len(section.header.paragraphs) > 0}')
    print(f'  - Has Footer: {len(section.footer.paragraphs) > 0}')
    print(f'  - Header paragraphs: {len(section.header.paragraphs)}')
    print(f'  - Footer paragraphs: {len(section.footer.paragraphs)}')
    if len(section.header.paragraphs) > 0:
        for j, h_para in enumerate(section.header.paragraphs):
            print(f'    Header {j}: {h_para.text[:50]}')
    if len(section.footer.paragraphs) > 0:
        for j, f_para in enumerate(section.footer.paragraphs):
            print(f'    Footer {j}: {f_para.text[:50]}')

print()
print('=== FIRST 10 PARAGRAPHS ===')
for i, para in enumerate(doc.paragraphs[:10]):
    text = (para.text or '')[:60]
    print(f'{i}: {repr(text)}')
