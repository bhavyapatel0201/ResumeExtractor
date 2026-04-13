#!/usr/bin/env python3
"""
Backend server to handle Word document generation with vetting report data.
This server provides endpoints to merge vetting report data into Word templates.
"""

import os
import json
import shutil
import tempfile
from datetime import datetime
from pathlib import Path
from http.server import HTTPServer, BaseHTTPRequestHandler
from urllib.parse import urlparse, parse_qs
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Try to import python-docx for Word document manipulation
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. Install with: pip install python-docx")


class VettingReportHandler(BaseHTTPRequestHandler):
    """HTTP request handler for vetting report document generation."""

    def do_OPTIONS(self):
        """Handle CORS preflight requests."""
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, GET, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()

    def do_GET(self):
        """Simple GET handler for health checks and root info."""
        parsed_path = urlparse(self.path)
        if parsed_path.path in ('/', '/health'):
            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            resp = json.dumps({'status': 'ok', 'path': parsed_path.path})
            self.wfile.write(resp.encode('utf-8'))
            return
        # Fallback to 404 for other GET paths
        self.send_response(404)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.end_headers()
        self.wfile.write(b'Not Found')

    def do_POST(self):
        """Handle POST requests to generate Word documents."""
        parsed_path = urlparse(self.path)
        
        if parsed_path.path == '/api/generate-report':
            self.handle_generate_report()
        elif parsed_path.path == '/api/generate-final-docx':
            self.handle_generate_final_docx()
        else:
            self.send_response(404)
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(b'Not Found')

    def handle_generate_report(self):
        """Generate a Word document with vetting report data."""
        try:
            # Read the request body
            content_length = int(self.headers.get('Content-Length', 0))
            body = self.rfile.read(content_length)
            payload = json.loads(body.decode('utf-8'))

            vetting_html = payload.get('vettingHtml', '')
            job_description = payload.get('jobDescription', '')
            resume_data = payload.get('resumeData', {})

            if not DOCX_AVAILABLE:
                self.send_response(500)
                self.send_header('Content-Type', 'application/json')
                self.send_header('Access-Control-Allow-Origin', '*')
                self.end_headers()
                error_response = json.dumps({'error': 'python-docx is not installed on the server.'})
                self.wfile.write(error_response.encode('utf-8'))
                return

            # Path to template file
            template_path = r'C:\Users\Mann\Resume Extractor\WORD FORMAT 2\Sample 2.docx'
            
            if not os.path.exists(template_path):
                self.send_response(404)
                self.send_header('Content-Type', 'application/json')
                self.send_header('Access-Control-Allow-Origin', '*')
                self.end_headers()
                error_response = json.dumps({'error': f'Template file not found: {template_path}'})
                self.wfile.write(error_response.encode('utf-8'))
                return

            # Create a temporary copy of the template
            temp_fd, temp_path = tempfile.mkstemp(suffix='.docx')
            os.close(temp_fd)
            shutil.copy(template_path, temp_path)

            # Open and modify the document
            doc = Document(temp_path)

            # Find a placeholder or add to the end
            # Look for a paragraph with "VETTING_REPORT_DATA" or add after first paragraph
            inserted = False
            for i, para in enumerate(doc.paragraphs):
                if 'VETTING_REPORT_DATA' in para.text or 'vetting' in para.text.lower():
                    # Replace this paragraph with vetting data
                    para.clear()
                    # No heading added here
                    self._add_html_content_to_doc(doc, i + 1, vetting_html, job_description)
                    inserted = True
                    break

            if not inserted:
                # Add vetting report at the end
                doc.add_paragraph()
                # No heading added here
                self._add_html_content_to_doc(doc, len(doc.paragraphs), vetting_html, job_description)

            # Save the modified document to bytes
            output_fd, output_path = tempfile.mkstemp(suffix='.docx')
            os.close(output_fd)
            doc.save(output_path)

            # Read the file and send as response
            with open(output_path, 'rb') as f:
                file_content = f.read()

            # Clean up temporary files
            os.remove(temp_path)
            os.remove(output_path)

            # Send response
            self.send_response(200)
            self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.send_header('Content-Disposition', 'attachment; filename="vetting-report.docx"')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.send_header('Content-Length', len(file_content))
            self.end_headers()
            self.wfile.write(file_content)

            logger.info('Successfully generated vetting report document')

        except Exception as e:
            logger.error(f'Error generating report: {str(e)}')
            self.send_response(500)
            self.send_header('Content-Type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            error_response = json.dumps({'error': f'Error generating report: {str(e)}'})
            self.wfile.write(error_response.encode('utf-8'))

    def handle_generate_final_docx(self):
        """Generate a combined Word document using SAMPLE final.docx template with outputs printed on it."""
        try:
            # Read the request body
            content_length = int(self.headers.get('Content-Length', 0))
            body = self.rfile.read(content_length)
            payload = json.loads(body.decode('utf-8'))

            output_html = payload.get('outputHtml', '')
            vetting_html = payload.get('vettingHtml', '')

            if not DOCX_AVAILABLE:
                self.send_response(500)
                self.send_header('Content-Type', 'application/json')
                self.send_header('Access-Control-Allow-Origin', '*')
                self.end_headers()
                error_response = json.dumps({'error': 'python-docx is not installed on the server.'})
                self.wfile.write(error_response.encode('utf-8'))
                return

            # Path to template file - use relative path from current working directory
            template_path = os.path.join(os.getcwd(), 'SAMPLE FINAL', 'FINAL TEMPLATE.docx')
            logger.info(f'Template path being used: {template_path}')
            
            if not os.path.exists(template_path):
                self.send_response(404)
                self.send_header('Content-Type', 'application/json')
                self.send_header('Access-Control-Allow-Origin', '*')
                self.end_headers()
                error_response = json.dumps({'error': f'Template file not found: {template_path}'})
                self.wfile.write(error_response.encode('utf-8'))
                return

            # Create a temporary copy of the template
            temp_fd, temp_path = tempfile.mkstemp(suffix='.docx')
            os.close(temp_fd)
            shutil.copy(template_path, temp_path)

            # Open the template document
            doc = Document(temp_path)

            # Preserve the template exactly and append vetting + output after it.
            # Do NOT remove or modify existing paragraphs/tables/sections — this keeps headers/footers intact.
            try:
                # Start vetting report on a new page
                doc.add_page_break()
            except Exception:
                # Fallback: insert explicit page break element
                from lxml import etree
                p = doc.add_paragraph()
                pPr = p._element.get_or_add_pPr()
                br = etree.SubElement(pPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br')
                br.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'page')

            # Append vetting report (preserves template header/footer)
            self._add_html_content_to_doc(doc, len(doc.paragraphs), vetting_html, '')

            # Insert a page break before OUTPUT
            try:
                doc.add_page_break()
            except Exception:
                from lxml import etree
                page_break_para = doc.add_paragraph()
                pPr = page_break_para._element.get_or_add_pPr()
                br = etree.SubElement(pPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br')
                br.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'page')

            # Append OUTPUT content
            self._add_html_content_to_doc_formatted(doc, len(doc.paragraphs), output_html, '')

            # Save the modified document to bytes
            output_fd, output_path = tempfile.mkstemp(suffix='.docx')
            os.close(output_fd)
            doc.save(output_path)

            # Read the file and send as response
            with open(output_path, 'rb') as f:
                file_content = f.read()

            # Clean up temporary files
            os.remove(temp_path)
            os.remove(output_path)

            # Send response
            self.send_response(200)
            self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.send_header('Content-Disposition', 'attachment; filename="final-output.docx"')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.send_header('Content-Length', len(file_content))
            self.end_headers()
            self.wfile.write(file_content)

            logger.info('Successfully generated final DOCX using template')

        except Exception as e:
            logger.error(f'Error generating final DOCX: {str(e)}')
            self.send_response(500)
            self.send_header('Content-Type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            error_response = json.dumps({'error': f'Error generating final DOCX: {str(e)}'})
            self.wfile.write(error_response.encode('utf-8'))

    def _add_html_content_to_doc(self, doc, insert_pos, html_content, job_description=None):
        """Add HTML content to the Word document with proper formatting."""
        try:
            import re
            text_content = re.sub(r'<[^>]+>', '\n', html_content)
            text_content = text_content.replace('&nbsp;', ' ')
            text_content = text_content.replace('&lt;', '<')
            text_content = text_content.replace('&gt;', '>')
            text_content = text_content.replace('&amp;', '&')

            lines = [line for line in text_content.split('\n') if line.strip()]
            # Remove any instruction lines that ask to analyze the job description
            # e.g. "Analyze the provided job description 'software engineer' against the resume and provide fit analysis."
            filtered = []
            for ln in lines:
                try:
                    # match variants containing 'analyze' + 'job description' and mentioning resume/against/provide fit
                    if re.search(r"analyze.*job description.*(resume|provide fit analysis|against)", ln, re.I):
                        # skip this instruction line when generating the Word document
                        continue
                    # also skip more generic patterns starting with 'analyze the job description'
                    if re.match(r"^\s*analyze the job description", ln, re.I):
                        continue
                except Exception:
                    pass
                filtered.append(ln)
            lines = filtered
            
            # Remove all leading empty lines
            while lines and not lines[0].strip():
                lines.pop(0)
            
            i = 0
            last_table = None
            def is_heading_token(s):
                if not s: return False
                t = s.lower().strip()
                tokens = [
                    'why is this resume fit for job description?',
                    'soft skills',
                    'skills assessment',
                    'work experience analysis',
                    'education verification',
                    'certifications review',
                    'projects evaluation',
                    'vetting summary',
                    'candidate information',
                    'job description',
                    'technical skills',
                    'technical'
                ]
                return tokens.index(t) >= 0 if t in tokens else False

            while i < len(lines):
                line = lines[i].strip()
                norm = line.lower().replace(' ', '')

                # Detect Technical Skills heading
                if line.lower().replace(' ', '') == 'technicalskills' or line.lower().strip() == 'technical skills' or line.lower().strip() == 'technical':
                    # Collect technical block until next major heading or soft skills
                    tech_block = []
                    i += 1
                    while i < len(lines):
                        nxt = lines[i].strip()
                        low = nxt.lower().strip()
                        if low.replace(' ', '') == 'softskills' or low in [
                            'why is this resume fit for job description?', 'skills assessment', 'work experience analysis', 'education verification', 'certifications review', 'projects evaluation', 'vetting summary', 'candidate information', 'job description', 'technical skills', 'technical'
                        ]:
                            break
                        tech_block.append(nxt)
                        i += 1

                    # Insert two blank paragraphs above the Technical Skills heading
                    doc.add_paragraph()
                    doc.add_paragraph()
                    # Prepare a 1x2 table for side-by-side skills
                    table = doc.add_table(rows=1, cols=2)
                    table.autofit = True
                    left_cell = table.cell(0,0)
                    right_cell = table.cell(0,1)

                    # Add Technical Skills heading
                    p = left_cell.paragraphs[0]
                    run = p.add_run('Technical Skills')
                    run.bold = True
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(0,0,0)

                    # Add technical block content
                    for item in tech_block:
                        if item.startswith('- '):
                            para = left_cell.add_paragraph(item[2:], style='List Bullet')
                            for r in para.runs:
                                r.font.size = Pt(12)
                        else:
                            para = left_cell.add_paragraph(item)
                            for r in para.runs:
                                r.font.size = Pt(12)

                    last_table = table
                    # do not increment i here, loop continues
                    continue

                # Detect Soft Skills heading
                if line.lower().replace(' ', '') == 'softskills' or line.lower().strip() == 'soft skills':
                    # Collect soft block
                    soft_block = []
                    i += 1
                    while i < len(lines):
                        nxt = lines[i].strip()
                        low = nxt.lower().strip()
                        if low in [
                            'why is this resume fit for job description?', 'skills assessment', 'work experience analysis', 'education verification', 'certifications review', 'projects evaluation', 'vetting summary', 'candidate information', 'job description', 'technical skills', 'technical'
                        ]:
                            break
                        soft_block.append(nxt)
                        i += 1

                    # If there is an existing skills table, fill its right cell. Otherwise create new table
                    if last_table is not None:
                        right_cell = last_table.cell(0,1)
                        p = right_cell.paragraphs[0]
                        run = p.add_run('Soft Skills')
                        run.bold = True
                        run.font.size = Pt(14)
                        run.font.color.rgb = RGBColor(0,0,0)
                        for item in soft_block:
                            if item.startswith('- '):
                                para = right_cell.add_paragraph(item[2:], style='List Bullet')
                                for r in para.runs:
                                    r.font.size = Pt(12)
                            else:
                                para = right_cell.add_paragraph(item)
                                for r in para.runs:
                                    r.font.size = Pt(12)
                    else:
                        # Insert two blank paragraphs above the Soft Skills heading/table
                        doc.add_paragraph()
                        doc.add_paragraph()
                        table = doc.add_table(rows=1, cols=2)
                        left_cell = table.cell(0,0)
                        right_cell = table.cell(0,1)
                        p = right_cell.paragraphs[0]
                        run = p.add_run('Soft Skills')
                        run.bold = True
                        run.font.size = Pt(14)
                        run.font.color.rgb = RGBColor(0,0,0)
                        for item in soft_block:
                            if item.startswith('- '):
                                para = right_cell.add_paragraph(item[2:], style='List Bullet')
                                for r in para.runs:
                                    r.font.size = Pt(12)
                            else:
                                para = right_cell.add_paragraph(item)
                                for r in para.runs:
                                    r.font.size = Pt(12)
                        last_table = table
                    continue

                # Generic heading detection
                if line.lower() in [
                    'why is this resume fit for job description?',
                    'skills assessment',
                    'work experience analysis',
                    'education verification',
                    'certifications review',
                    'projects evaluation',
                    'vetting summary',
                    'candidate information',
                    'job description'
                ]:
                    para = doc.add_paragraph()
                    run = para.add_run(line)
                    run.bold = True
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # (Previously insertion of instruction text was here; removed to keep DOCX clean)
                    i += 1
                    continue

                # Bullet point detection
                if line.startswith('- '):
                    para = doc.add_paragraph(line[2:], style='List Bullet')
                    for run in para.runs:
                        run.font.size = Pt(12)
                        run.bold = False
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    i += 1
                    continue

                # Remove duplicate lowercase question line from Word output
                if line.strip().lower() == 'why is this resume fit for job description:':
                    i += 1
                    continue
                para = doc.add_paragraph(line)
                for run in para.runs:
                    run.font.size = Pt(12)
                    run.bold = False
                    run.font.color.rgb = RGBColor(0, 0, 0)
                i += 1
        except Exception as e:
            logger.error(f'Error adding HTML content: {str(e)}')
            para = doc.add_paragraph(f'Error processing content: {str(e)}')
            for run in para.runs:
                run.font.size = Pt(12)
                run.bold = False
                run.font.color.rgb = RGBColor(0, 0, 0)


    def _add_html_content_to_doc_formatted(self, doc, insert_pos, html_content, job_description=None):
        """Add HTML content to the Word document matching 'Download OUTPUT' format: 14pt bold headings with spacing, 12pt data."""
        try:
            import re
            text_content = re.sub(r'<[^>]+>', '\n', html_content)
            text_content = text_content.replace('&nbsp;', ' ')
            text_content = text_content.replace('&lt;', '<')
            text_content = text_content.replace('&gt;', '>')
            text_content = text_content.replace('&amp;', '&')

            lines = [line for line in text_content.split('\n') if line.strip()]
            
            # Filter out instruction lines
            filtered = []
            for ln in lines:
                try:
                    if re.search(r"analyze.*job description.*(resume|provide fit analysis|against)", ln, re.I):
                        continue
                    if re.match(r"^\s*analyze the job description", ln, re.I):
                        continue
                except Exception:
                    pass
                filtered.append(ln)
            lines = filtered
            
            # Remove all leading empty lines
            while lines and not lines[0].strip():
                lines.pop(0)

            # Enhanced processing: support side-by-side Technical Skills / Soft Skills table
            heading_list = [
                'candidate information',
                'job description',
                'technical skills',
                'soft skills',
                'work experience',
                'education',
                'certifications',
                'projects',
                'summary',
                'analysis',
                'recommendations',
                'fit analysis',
                'skills match',
                'skills'  # Add "skills" as a heading
            ]

            i = 0
            last_table = None
            while i < len(lines):
                line = lines[i].strip()
                if not line:
                    i += 1
                    continue

                low = line.lower().strip()

                # Handle Technical Skills block
                if low in ('technical skills', 'technical'):
                    tech_block = []
                    j = i + 1
                    while j < len(lines):
                        nxt = lines[j].strip()
                        if nxt.lower().strip() in heading_list:
                            break
                        tech_block.append(nxt)
                        j += 1

                    # Create or reuse a 1x2 table for skills
                    if last_table is None:
                        table = doc.add_table(rows=1, cols=2)
                        table.autofit = True
                        last_table = table
                    else:
                        table = last_table

                    left_cell = table.cell(0, 0)
                    # Clear any existing default paragraph
                    left_cell.paragraphs[0].clear()
                    run = left_cell.paragraphs[0].add_run('Technical Skills')
                    run.bold = True
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(0, 0, 0)

                    for item in tech_block:
                        if item.startswith('- '):
                            para = left_cell.add_paragraph(item[2:], style='List Bullet')
                        else:
                            para = left_cell.add_paragraph(item)
                        for r in para.runs:
                            r.font.size = Pt(12)

                    i = j
                    continue

                # Handle Soft Skills block
                if low in ('soft skills', 'soft'):
                    soft_block = []
                    j = i + 1
                    while j < len(lines):
                        nxt = lines[j].strip()
                        if nxt.lower().strip() in heading_list:
                            break
                        soft_block.append(nxt)
                        j += 1

                    if last_table is None:
                        table = doc.add_table(rows=1, cols=2)
                        table.autofit = True
                        last_table = table
                    else:
                        table = last_table

                    right_cell = table.cell(0, 1)
                    right_cell.paragraphs[0].clear()
                    run = right_cell.paragraphs[0].add_run('Soft Skills')
                    run.bold = True
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(0, 0, 0)

                    for item in soft_block:
                        if item.startswith('- '):
                            para = right_cell.add_paragraph(item[2:], style='List Bullet')
                        else:
                            para = right_cell.add_paragraph(item)
                        for r in para.runs:
                            r.font.size = Pt(12)

                    i = j
                    continue

                # Generic heading detection
                is_heading = low in heading_list
                if is_heading:
                    para = doc.add_paragraph()
                    run = para.add_run(line)
                    run.bold = True
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    i += 1
                    continue

                # Bullet points and regular lines
                if line.startswith('- '):
                    para = doc.add_paragraph(style='List Bullet')
                    bullet_text = line[2:]
                    run = para.add_run(bullet_text)
                    run.font.size = Pt(12)
                    run.bold = False
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    i += 1
                    continue

                # Regular data line
                para = doc.add_paragraph(line)
                for run in para.runs:
                    run.font.size = Pt(12)
                    run.bold = False
                    run.font.color.rgb = RGBColor(0, 0, 0)
                i += 1
        except Exception as e:
            logger.error(f'Error adding formatted HTML content: {str(e)}')
            para = doc.add_paragraph(f'Error processing content: {str(e)}')
            for run in para.runs:
                run.font.size = Pt(12)
                run.bold = False
                run.font.color.rgb = RGBColor(0, 0, 0)

    def _remove_leading_empty_paragraphs(self, doc):
        """Remove all leading empty paragraphs from the document."""
        try:
            while len(doc.paragraphs) > 0:
                para = doc.paragraphs[0]
                text = (para.text or '').strip()
                
                # Check if paragraph has any visible content
                if not text:
                    # Check if paragraph has any runs with content
                    has_content = False
                    for run in para.runs:
                        if (run.text or '').strip():
                            has_content = True
                            break
                    
                    if not has_content:
                        # Remove this empty paragraph
                        p = para._element
                        parent = p.getparent()
                        if parent is not None:
                            parent.remove(p)
                        continue
                
                # Found first non-empty paragraph, stop
                break
        except Exception as e:
            logger.debug(f'Error removing leading empty paragraphs: {e}')

    def log_message(self, format, *args):
        """Override to use logger instead of stderr."""
        logger.info("%s - - [%s] %s" % (
            self.client_address[0],
            self.log_date_time_string(),
            format % args))


def run_server(port=8443, host='0.0.0.0'):
    """Run the backend server."""
    server_address = (host, port)
    httpd = HTTPServer(server_address, VettingReportHandler)
    logger.info(f'Backend server running at http://{host}:{port}')
    try:
        httpd.serve_forever()
    except KeyboardInterrupt:
        logger.info('Server stopped.')
        httpd.shutdown()


if __name__ == '__main__':
    # Use 0.0.0.0 so local browser (file://) and other interfaces can reach the server
    run_server(8444, '0.0.0.0')
