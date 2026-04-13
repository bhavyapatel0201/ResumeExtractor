from docx import Document
from lxml import etree
import os

def inspect(path):
    print(f"\n=== Inspecting: {path} ===")
    if not os.path.exists(path):
        print('MISSING')
        return
    doc = Document(path)
    print(f"paragraphs: {len(doc.paragraphs)}, tables: {len(doc.tables)}, sections: {len(doc.sections)}")
    for i, sec in enumerate(doc.sections):
        print(f"-- Section {i} --")
        hdr_text = '\n'.join([p.text for p in sec.header.paragraphs if p.text.strip()]) or '<empty>'
        ftr_text = '\n'.join([p.text for p in sec.footer.paragraphs if p.text.strip()]) or '<empty>'
        print('header:', hdr_text)
        print('footer:', ftr_text)
        try:
            xml = etree.tostring(sec._sectPr, pretty_print=True, encoding='unicode')
            print('sectPr snippet:', xml.splitlines()[0])
        except Exception as e:
            print('sectPr dump error', e)


template = r"C:\Users\Mann\Resume Extractor\SAMPLE FINAL\FINAL TEMPLATE.docx"
generated = r"C:\Users\Mann\Resume Extractor\tmp_test.docx"
inspect(template)
inspect(generated)
print('\nDone')
