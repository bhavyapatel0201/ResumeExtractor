from docx import Document
import os

template = r"C:\Users\Mann\Resume Extractor\SAMPLE FINAL\FINAL TEMPLATE.docx"
out = r"C:\Users\Mann\Resume Extractor\tmp_local.docx"
if not os.path.exists(template):
    print('template missing')
    raise SystemExit(1)

doc = Document(template)
print('Sections before:', len(doc.sections))
# Append vetting
try:
    doc.add_page_break()
except Exception:
    pass
p = doc.add_paragraph('VETTING TEST: This is a test')
# Append output
try:
    doc.add_page_break()
except Exception:
    pass
p2 = doc.add_paragraph('OUTPUT TEST: This is a test')

doc.save(out)
print('Saved', out)
