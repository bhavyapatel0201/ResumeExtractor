from docx import Document
from lxml import etree
import os

path = r"C:\Users\Mann\Resume Extractor\SAMPLE FINAL\FINAL TEMPLATE.docx"
print(f"Template path: {path}")
if not os.path.exists(path):
    print("File not found")
    raise SystemExit(1)

doc = Document(path)
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")
print(f"Total sections: {len(doc.sections)}")

for i, sec in enumerate(doc.sections):
    print('\n--- Section', i, '---')
    try:
        pg_size = (sec.page_height, sec.page_width)
    except Exception:
        pg_size = None
    print('page_size:', pg_size)
    # Header
    hdr = sec.header
    hdr_text = '\n'.join([p.text for p in hdr.paragraphs if p.text.strip()])
    print('header paragraphs:', len(hdr.paragraphs))
    if hdr_text:
        print('header text:')
        print(hdr_text)
    else:
        print('header text: <empty>')
    # Footer
    ftr = sec.footer
    ftr_text = '\n'.join([p.text for p in ftr.paragraphs if p.text.strip()])
    print('footer paragraphs:', len(ftr.paragraphs))
    if ftr_text:
        print('footer text:')
        print(ftr_text)
    else:
        print('footer text: <empty>')
    # Section XML
    try:
        xml = etree.tostring(sec._sectPr, pretty_print=True, encoding='unicode')
        print('sectPr xml snippet:')
        print('\n'.join(xml.splitlines()[:20]))
    except Exception as e:
        print('could not dump sectPr xml:', e)

# Print first 40 paragraphs text
print('\n--- First 40 document paragraphs ---')
for i, p in enumerate(doc.paragraphs[:40]):
    txt = p.text.replace('\n', ' ')
    print(f'[{i}] ({len(p.runs)} runs) "{txt}"')

# Check for explicit page breaks in runs
print('\n--- Paragraphs with page breaks in runs ---')
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        if '\n' in run.text:
            pass
    # raw xml check for br elements
    try:
        xml = etree.tostring(p._element, encoding='unicode')
        if '<w:br' in xml or 'type="page"' in xml:
            print(f'Paragraph {i} has break element in xml')
    except Exception:
        pass

print('\n--- Done ---')
